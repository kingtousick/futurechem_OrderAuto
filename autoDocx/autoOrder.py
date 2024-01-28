import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
import sqlite3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime
from copy import deepcopy
import os
import ctypes
import sys 
from tkinter import messagebox
# pyinstaller -w -F --icon=C:/Users/wjdck/OneDrive/문서/futurechem_OrderAuto/futurechem_OrderAuto/autoDocx/icon/futureMain2.ico  C:/Users/wjdck/OneDrive/문서/futurechem_OrderAuto/futurechem_OrderAuto/autoDocx/autoOrder.py

import ctypes

# 변경할 아이콘 파일 경로
icon_path = 'C:/Users/wjdck/OneDrive/문서/futurechem_OrderAuto/futurechem_OrderAuto/autoDocx/icon/futureMain2.ico'

# 아이콘 적용
ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(icon_path)

# 창 제목 표시줄 아이콘 변경
ctypes.windll.user32.SendMessageW(ctypes.windll.kernel32.GetConsoleWindow(), 0x80, 0, icon_path)


def resource_path(relative_path):
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


class PurchaseRequestApp:
    def __init__(self, root):

        # 추가 코드: 창 아이콘 변경
        root.iconbitmap(default=icon_path)

        self.root = root
        self.root.title("구매요구서 자동화")
        self.manpower_entries = [] 

        self.tab_control = ttk.Notebook(root)
        
        # Tab 1: Manual Input
        self.tab1 = ttk.Frame(self.tab_control)
        self.tab_control.add(self.tab1, text="수동입력")
        self.setup_tab1()

        # Tab 2: Data Registration
        self.tab2 = ttk.Frame(self.tab_control)
        self.tab_control.add(self.tab2, text="데이터 등록")
        self.setup_tab2()

        # Tab 3: Search and Update
        self.tab3 = ttk.Frame(self.tab_control)
        self.tab_control.add(self.tab3, text="검색 입력")
        self.setup_tab3()

        self.tab_control.pack(expand=1, fill='both')
      
    def setup_tab1(self):
    
       label = ttk.Label(self.tab1, text="구매요구서 양식:")
       label.grid(column=0, row=0)

       self.file_path = tk.StringVar()

       entry = ttk.Entry(self.tab1, textvariable=self.file_path, state="readonly")
       entry.grid(column=1, row=0)

       button = ttk.Button(self.tab1, text="Browse", command=self.upload_form)
       button.grid(column=2, row=0)

       # 사용자 입력 부분 
       ttk.Label(self.tab1, text="기관명:").grid(column=0, row=1, sticky=tk.E)
       self.institution_entry = ttk.Entry(self.tab1)
       self.institution_entry.grid(column=1, row=1)
       
       ttk.Label(self.tab1, text="기관코드:").grid(column=0, row=2, sticky=tk.E)
       self.institution_code_entry = ttk.Entry(self.tab1)
       self.institution_code_entry.grid(column=1, row=2)

       ttk.Label(self.tab1, text="핵종명[앞]:").grid(column=0, row=3, sticky=tk.E)
       self.isotope_fname_entry = ttk.Entry(self.tab1)
       self.isotope_fname_entry.grid(column=1, row=3)

       ttk.Label(self.tab1, text="핵종수량:").grid(column=0, row=4, sticky=tk.E)
       self.quantity_entry = ttk.Entry(self.tab1)
       self.quantity_entry.grid(column=1, row=4)

       
       ttk.Label(self.tab1, text="핵종명[뒤]:").grid(column=0, row=5, sticky=tk.E)
       self.isotope_ename_entry = ttk.Entry(self.tab1)
       self.isotope_ename_entry.grid(column=1, row=5)

       ttk.Label(self.tab1, text="모델명:").grid(column=0, row=6, sticky=tk.E)
       self.model_name_entry = ttk.Entry(self.tab1)
       self.model_name_entry.grid(column=1, row=6)
       
     
       manpower_labels = [f"투입인력 : {i:02d}:{j:02d}" for i in range(8, 18) for j in range(0, 60, 30) if not ((i == 8 and j == 0) or (i == 17 and j == 30))]
       for i in range(len(manpower_labels)):
            ttk.Label(self.tab1, text=manpower_labels[i]).grid(column=0, row=7+i)
            entry = ttk.Entry(self.tab1)
            entry.grid(column=1, row=7+i)
            self.manpower_entries.append(entry)
       
       ttk.Label(self.tab1, text="구입 예정일:").grid(column=0, row=25, sticky=tk.E)
       self.purchase_date_entry = ttk.Entry(self.tab1)
       self.purchase_date_entry.grid(column=1, row=25)

       # 워드파일 생성
       generate_button = ttk.Button(self.tab1, text="파일생성하기", command=self.generate_word_file)
       generate_button.grid(column=1, row=27)

    def upload_form(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if file_path:
            self.file_path.set(file_path)
                                              
    def generate_word_file(self):
        template_path = self.file_path.get()
        if template_path:
            # 양식 파일 
            # original_document = Document(template_path)
            document = Document(template_path)
        
            # 문서를 복사하여 기존 양식 유지
            # document = deepcopy(original_document)

            # 파일내용 변환 
            placeholders = {
                "instNm": self.institution_entry.get(),
                "instCd": self.institution_code_entry.get(),
                "isotopeFront": self.isotope_fname_entry.get(),
                "Quantity": self.quantity_entry.get (),
                "isotopeEnd": self.isotope_ename_entry.get(),
                "model": self.model_name_entry.get(),
                "purchDt": self.purchase_date_entry.get()
            }
            # manPw 엔트리들을 동적으로 추가
            for i in range(1, 19):
              manPw_value = self.manpower_entries[i-1].get()
              if not manPw_value:  # 공백인 경우
               placeholders[f"manPw{i}"] = self.manpower_entries[i-1].insert(0, " ")  # 스페이스바 추가
              else: 
                placeholders[f"manPw{i}"] = self.manpower_entries[i-1].get()
                print(f"manPw{i}={self.manpower_entries[i-1].get()}")
            
            # 각 표에서 검색 단어를 치환합니다.
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        print(f"CELL={cell.text}")
                        for key, value in placeholders.items():
                            if key in cell.text:
                                import re
                                cell.text = re.sub(rf"\b{re.escape(key)}\b", value, cell.text)                       
                                print(f"After replace - {key}={value}, cell.text={cell.text}")            
                           
                            # # 텍스트 정렬을 가운데 정렬로 변경 
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                #  텍스트 정렬을 가운데 정렬로 변경
                                # for paragraph in cell.paragraphs:
                                #     for run in paragraph.runs:
                                #         run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                                # 병합된 셀의 경우 수직 정렬도 설정
                                if cell._element.xpath(".//w:vMerge") and cell._element.xpath(".//w:vMerge")[0].values()[0] == "restart":
                                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                
                                    

                                # 높이 설정 (예시에서는 줄 수에 따라 조절)
                                # lines = len(value.split('\n'))
                                # cell.paragraphs[0].runs[0].font.size = Pt(12)  # 원하는 폰트 크기로 설정
                                # cell.height = Inches(lines * 0.3)  # 적절한 상수를 곱하여 높이 설정    
     

            # 변경된 파일 저장
            now = datetime.now()
            timestamp = now.strftime("%Y%m%d_%H%M%S")
            default_file_name = f"구매요구서_{timestamp}.docx"

            # 사용자에게 저장할 파일의 경로 선택 받기
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                initialfile=default_file_name
            )

            if file_path:
                document.save(file_path)
                print(f"파일이 저장되었습니다: {file_path}")

            # "완료 되었습니다" 얼럿창 띄우기
            messagebox.showinfo("완료", "파일 생성이 완료되었습니다!")
        
        pass

    def setup_tab2(self):
        # Implement Tab 2 UI and functionality here       
        label = ttk.Label(self.tab2, text="■■■■■■■■■■■기능 미정■■■■■■■■■■■")
        label.grid(column=0, row=0)
        pass

    def setup_tab3(self):
        # Implement Tab 3 UI and functionality here
        label = ttk.Label(self.tab3, text="■■■■■■■■■■■기능 미정■■■■■■■■■■■")
        label.grid(column=0, row=0)
        pass

def main():
    root = tk.Tk()
    app = PurchaseRequestApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
